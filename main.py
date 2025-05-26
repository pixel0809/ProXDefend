from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify
import psutil
import os
import hashlib
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from datetime import datetime
import io
import time
import magic
import requests
from dotenv import load_dotenv
from cache_manager import CacheManager
import json
import socket
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import traceback
import PyPDF2
from docx import Document
import re
import math
import yara
import struct
import binascii
import subprocess
import ipaddress

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max file size
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize cache manager
cache_manager = CacheManager(cache_dir="cache", cache_duration=21600)  # 6 hours in seconds

# Register fonts
try:
    # Try to register Helvetica font with a more robust path check
    font_path = os.path.join(os.path.dirname(__file__), 'fonts', 'Helvetica.ttf')
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont('Helvetica', font_path))
    else:
        # Fallback to default font if Helvetica is not available
        pdfmetrics.registerFont(TTFont('DefaultFont', 'Arial'))
except Exception as e:
    print(f"Warning: Using default font - {str(e)}")

# Suspicious keywords for local scanning
SUSPICIOUS_KEYWORDS = ["powershell", "cmd", "nc", "ncat", "mimikatz", "evil", "hack", "exploit", "payload", "trojan", "malware",
                     "sh -i", "/dev/tcp", "reverse shell", "bind shell", "netcat", "socat", "bash -i", ">& /dev/tcp",
                     "python -c", "perl -e", "ruby -rsocket", "php -r", "fsockopen", "socket.connect", "exec('/bin/bash')",
                     "system('/bin/bash')", "ProcessBuilder", "pty.spawn", "/bin/sh", "0>&1"]

# List of suspicious ports
SUSPICIOUS_PORTS = [
    # Common malware/backdoor ports
    4444, 1337, 6666, 12345, 5555,
    # Remote access ports
    22, 23, 3389, 5900, 5800, 5000,
    # Database ports
    1433, 1434, 3306, 5432, 27017,
    # Web server ports
    80, 443, 8080, 8443,
    # File transfer ports
    21, 20, 69,
    # Mail ports
    25, 465, 587,
    # DNS ports
    53,
    # RPC ports
    135, 139, 445,
    # Other potentially dangerous ports
    1080, 3128, 8081, 8443, 8888
]

# Add after SUSPICIOUS_PORTS
SUSPICIOUS_IP_RANGES = [
    '192.168.0.0/16',  # Private network
    '10.0.0.0/8',      # Private network
    '172.16.0.0/12',   # Private network
    '169.254.0.0/16',  # Link-local
    '127.0.0.0/8'      # Localhost
]

# Add after the SUSPICIOUS_KEYWORDS definition
WINDOWS_SYSTEM_PROCESSES = {
    'svchost.exe',      # Windows Service Host
    'lsass.exe',        # Windows Security
    'csrss.exe',        # Windows Subsystem
    'winlogon.exe',     # Windows Logon
    'services.exe',     # Windows Services
    'smss.exe',         # Windows Session Manager
    'explorer.exe',     # Windows Explorer
    'spoolsv.exe',      # Print Spooler
    'wininit.exe',      # Windows Start-Up
    'fontdrvhost.exe',  # Font Driver Host
    'dwm.exe',          # Desktop Window Manager
    'taskmgr.exe',      # Task Manager
    'RuntimeBroker.exe', # Runtime Broker
    'WUDFHost.exe',     # Windows Driver Foundation Host
    'cmd.exe',          # Command Prompt
    'msedgewebview2.exe', # Microsoft Edge WebView2
    'powershell.exe'    # Windows PowerShell
}

# Add after WINDOWS_SYSTEM_PROCESSES
CRITICAL_SYSTEM_PROCESSES = {
    'System',
    'Registry',
    'smss.exe',
    'csrss.exe',
    'wininit.exe',
    'services.exe',
    'lsass.exe',
    'winlogon.exe',
    'explorer.exe',
    'svchost.exe'
}

# Dictionary of fix suggestions for each suspicious keyword
KEYWORD_FIX_SUGGESTIONS = {
    "powershell": "Check for malicious scripts. Disable unnecessary PowerShell access if possible.",
    "cmd": "Investigate command-line usage. Restrict cmd.exe usage via policies if abused.",
    "nc": "Possible reverse shell (netcat). Kill process immediately and inspect the system for backdoors.",
    "ncat": "Detected Netcat alternative. Analyze open ports and active sessions immediately.",
    "mimikatz": "Password dumping tool detected. Change all passwords and perform full forensic investigation.",
    "evil": "Potential exploitation tool. Terminate process and isolate the system from the network.",
    "hack": "Possible hacking attempt. Conduct a full system integrity check.",
    "exploit": "Exploit tool detected. Patch system vulnerabilities immediately.",
    "payload": "Malware payload delivery. Disconnect from the network and perform a full scan.",
    "trojan": "Trojan horse behavior. Remove process, quarantine system, and scan deeply.",
    "malware": "Generic malware detected. Isolate the device and run antivirus plus manual checks."
}

# Advanced scanning patterns
MALWARE_PATTERNS = {
    'shellcode': rb'\x90{20,}|\xCC{20,}',  # NOP sleds or INT3 instructions
    'powershell_encoded': rb'powershell.*-enc\s+[A-Za-z0-9+/=]+',
    'base64_encoded': rb'[A-Za-z0-9+/]{50,}={0,2}',
    'hex_encoded': rb'0x[0-9A-Fa-f]{50,}',
    'suspicious_commands': rb'(cmd\.exe|powershell\.exe|wscript\.exe|mshta\.exe).*\/c\s+',
    'network_indicators': rb'(https?:\/\/[^\s<>"]+|www\.[^\s<>"]+)',
    'registry_modifications': rb'reg\s+(add|delete|modify)',
    'file_operations': rb'(copy|move|del|rmdir)\s+[^\s]+\s+[^\s]+',
    'suspicious_functions': rb'(CreateRemoteThread|VirtualAlloc|WriteProcessMemory|ShellExecute)',
    'obfuscated_strings': rb'[A-Za-z0-9+/]{20,}={0,2}',
    'suspicious_ips': rb'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    'suspicious_ports': rb':(4444|1337|6666|12345|5555)',
    'suspicious_extensions': rb'\.(exe|dll|bat|ps1|vbs|js|hta|scr|pif|cmd|reg|inf|lnk)$',
    'suspicious_macros': rb'(VBA|macro|AutoOpen|Document_Open)',
    'suspicious_scripts': rb'(eval|exec|system|subprocess|os\.system)',
    'suspicious_network': rb'(socket|connect|bind|listen|accept)',
    'suspicious_crypto': rb'(crypto|encrypt|decrypt|AES|RSA|MD5|SHA)',
    'suspicious_antidebug': rb'(IsDebuggerPresent|CheckRemoteDebuggerPresent|OutputDebugString)',
    'suspicious_persistence': rb'(Startup|RunOnce|RunServices|Winlogon)',
    'reverse_shell': rb'(sh\s*-i\s*>[>&]\s*/dev/tcp/[0-9.]+/\d+|nc\s+-e\s+/bin/[a-z]+|bash\s*-i\s*>[>&]\s*/dev/tcp/|python\s+-c\s*[\'"]import\s+socket|perl\s+-e\s*[\'"]use\s+Socket|ruby\s+-rsocket|php\s+-r\s*[\'"].*fsockopen|socat\s+.*tcp-connect)'
}

# YARA rules for malware detection
YARA_RULES = {
    'suspicious_macro': '''
        rule SuspiciousMacro {
            meta:
                description = "Detects suspicious VBA macros"
            strings:
                $auto_open = "AutoOpen" nocase
                $document_open = "Document_Open" nocase
                $shell = "Shell" nocase
                $wscript = "WScript" nocase
                $powershell = "PowerShell" nocase
            condition:
                2 of them
        }
    ''',
    'suspicious_pdf': '''
        rule SuspiciousPDF {
            meta:
                description = "Detects suspicious PDF content"
            strings:
                $js = "/JS" nocase
                $javascript = "/JavaScript" nocase
                $launch = "/Launch" nocase
                $open_action = "/OpenAction" nocase
                $exec = "/Exec" nocase
                $cmd = "/Cmd" nocase
                $powershell = "/PowerShell" nocase
            condition:
                3 of them
        }
    ''',
    'suspicious_script': '''
        rule SuspiciousScript {
            meta:
                description = "Detects suspicious script content"
            strings:
                $eval = "eval(" nocase
                $exec = "exec(" nocase
                $system = "system(" nocase
                $shell = "shell(" nocase
                $cmd = "cmd.exe" nocase
                $powershell = "powershell.exe" nocase
            condition:
                2 of them
        }
    '''
}

# Function to calculate SHA256 of a file
def get_file_hash(file_path):
    try:
        with open(file_path, "rb") as f:
            return hashlib.sha256(f.read()).hexdigest()
    except Exception:
        return None

# Updated Basic File Scanning Function with Keyword Detection
def scan_file_basic(file_path):
    try:
        print(f"Starting basic scan for file: {file_path}")

        # Check file size (limit to 5MB for scanning)
        file_size = os.path.getsize(file_path)
        if file_size > 5 * 1024 * 1024:
            print(f"File too large to scan easily: {file_size} bytes")
            return ("Unknown (Too Large)", [])

        # Detect file type
        file_type = magic.from_file(file_path, mime=True)
        print(f"Detected file type: {file_type}")

        found_keywords = []

        # Handle different file types
        if file_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            try:
                doc = Document(file_path)
                # Check text in paragraphs
                for paragraph in doc.paragraphs:
                    text = paragraph.text.lower()
                    for keyword in SUSPICIOUS_KEYWORDS:
                        if keyword.lower() in text:
                            found_keywords.append(keyword)

                # Check text in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.lower()
                            for keyword in SUSPICIOUS_KEYWORDS:
                                if keyword.lower() in text:
                                    found_keywords.append(keyword)

                # Check text in headers and footers
                for section in doc.sections:
                    if section.header:
                        text = section.header.text.lower()
                        for keyword in SUSPICIOUS_KEYWORDS:
                            if keyword.lower() in text:
                                found_keywords.append(keyword)
                    if section.footer:
                        text = section.footer.text.lower()
                        for keyword in SUSPICIOUS_KEYWORDS:
                            if keyword.lower() in text:
                                found_keywords.append(keyword)

            except Exception as e:
                print(f"Error scanning Word document: {e}")
                return ("Error (Word Document)", [])

        elif file_type == 'application/pdf':
            try:
                pdf = PyPDF2.PdfReader(file_path)
                # Extract and scan text from all pages
                for page in pdf.pages:
                    text = page.extract_text().lower()
                    for keyword in SUSPICIOUS_KEYWORDS:
                        if keyword.lower() in text:
                            found_keywords.append(keyword)

            except Exception as e:
                print(f"Error scanning PDF: {e}")
                return ("Error (PDF)", [])

        else:
            # Handle text-based files
            text_based_types = ['text/', 'application/json', 'application/xml', 'application/javascript', 
                              'application/x-httpd-php', 'application/x-sh', 'application/x-python']
            
            if any(file_type.startswith(t) for t in text_based_types):
                try:
                    with open(file_path, 'r', errors='ignore') as f:
                        content = f.read().lower()
                        for keyword in SUSPICIOUS_KEYWORDS:
                            if keyword in content:
                                found_keywords.append(keyword)
                except Exception as e:
                    print(f"Error reading text file: {e}")
                    return ("Error (Text File)", [])
            else:
                # For binary files, try reading as bytes and search for string patterns
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                        # Convert keywords to bytes for binary search
                        for keyword in SUSPICIOUS_KEYWORDS:
                            keyword_bytes = keyword.encode('utf-8', errors='ignore')
                            if keyword_bytes in content:
                                found_keywords.append(keyword)
                except Exception as e:
                    print(f"Error reading binary file: {e}")
                    return ("Error (Binary File)", [])

        if found_keywords:
            print(f"Suspicious keywords found: {found_keywords}")
            return ("Malicious (Basic Scan)", list(set(found_keywords)))  # Remove duplicates
        else:
            return ("Clean (Basic Scan)", [])

    except Exception as e:
        print("Error scanning file:", e)
        return ("Error", [])

def calculate_entropy(data):
    """Calculate Shannon entropy of data."""
    if not data:
        return 0
    entropy = 0
    for x in range(256):
        p_x = data.count(bytes([x])) / len(data)
        if p_x > 0:
            entropy += -p_x * math.log2(p_x)
    return entropy

def analyze_file_structure(file_path):
    """Analyze file structure for suspicious patterns."""
    structure_analysis = {
        'suspicious_sections': [],
        'entropy_analysis': [],
        'file_header': None,
        'file_footer': None
    }
    
    try:
        with open(file_path, 'rb') as f:
            # Read file header and footer
            header = f.read(16)
            f.seek(-16, 2)
            footer = f.read()
            
            structure_analysis['file_header'] = binascii.hexlify(header).decode()
            structure_analysis['file_footer'] = binascii.hexlify(footer).decode()
            
            # Calculate entropy for different sections
            f.seek(0)
            chunk_size = 1024
            while True:
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                entropy = calculate_entropy(chunk)
                if entropy > 7.0:  # High entropy might indicate encryption or obfuscation
                    structure_analysis['entropy_analysis'].append({
                        'offset': f.tell() - len(chunk),
                        'entropy': entropy
                    })
    except Exception as e:
        print(f"Error analyzing file structure: {e}")
    
    return structure_analysis

def scan_file_enhanced(file_path):
    try:
        # Get file type
        file_type = magic.from_file(file_path, mime=True)
        
        # Initialize scan results
        found_patterns = []
        scan_details = []
        
        # Calculate file hash
        file_hash = get_file_hash(file_path)
        
        # Read file content
        with open(file_path, 'rb') as f:
            content = f.read()
            
            # Calculate overall file entropy
            file_entropy = calculate_entropy(content)
            
            # Adjust entropy thresholds based on file type
            if file_type == 'application/pdf':
                # PDFs naturally have higher entropy due to their binary format
                entropy_threshold = 8.0  # Increased threshold for PDFs
                if file_entropy > entropy_threshold:
                    scan_details.append(f"High overall file entropy ({file_entropy:.2f}) - possible encryption or obfuscation")
                    found_patterns.append('high_entropy')
            else:
                if file_entropy > 7.0:
                    scan_details.append(f"High overall file entropy ({file_entropy:.2f}) - possible encryption or obfuscation")
                    found_patterns.append('high_entropy')
            
            # Apply YARA rules
            for rule_name, rule_content in YARA_RULES.items():
                try:
                    rules = yara.compile(source=rule_content)
                    matches = rules.match(data=content)
                    if matches:
                        found_patterns.append(f"yara_{rule_name}")
                        scan_details.append(f"Matched YARA rule: {rule_name}")
                except Exception as e:
                    scan_details.append(f"Error applying YARA rule {rule_name}: {str(e)}")
            
            # Check for malware patterns with more specific criteria
            for pattern_name, pattern in MALWARE_PATTERNS.items():
                if pattern_name in ['base64_encoded', 'obfuscated_strings', 'suspicious_crypto', 'network_indicators', 'suspicious_scripts']:
                    # Skip these patterns for PDFs as they often cause false positives
                    if file_type == 'application/pdf':
                        continue
                if re.search(pattern, content, re.IGNORECASE):
                    found_patterns.append(pattern_name)
                    scan_details.append(f"Found {pattern_name} pattern")
        
        # Handle different file types
        if file_type == 'application/pdf':
            try:
                pdf = PyPDF2.PdfReader(file_path)
                
                # Check for JavaScript with more specific criteria
                has_js = False
                has_suspicious_js = False
                
                for page in pdf.pages:
                    try:
                        page_obj = page.get_object()
                        if isinstance(page_obj, dict) and '/JS' in page_obj:
                            has_js = True
                            # Check for suspicious JavaScript content
                            js_content = page_obj['/JS']
                            if isinstance(js_content, str) and any(keyword in js_content.lower() for keyword in ['eval', 'exec', 'system', 'shell', 'cmd', 'powershell']):
                                has_suspicious_js = True
                                found_patterns.append('pdf_suspicious_js')
                                scan_details.append("PDF contains suspicious JavaScript")
                    except Exception:
                        continue
                
                if has_js and not has_suspicious_js:
                    # JavaScript present but not suspicious
                    scan_details.append("PDF contains JavaScript (non-suspicious)")
                
                # Check for embedded files with size limit
                try:
                    if pdf.trailer and isinstance(pdf.trailer, dict):
                        root = pdf.trailer.get('/Root', {})
                        if isinstance(root, dict):
                            names = root.get('/Names', {})
                            if isinstance(names, dict):
                                embedded_files = names.get('/EmbeddedFiles', [])
                                if isinstance(embedded_files, list) and len(embedded_files) > 5:
                                    found_patterns.append('pdf_embedded_files')
                                    scan_details.append("PDF contains multiple embedded files")
                except Exception:
                    pass
                
                # Check for suspicious open actions
                try:
                    if pdf.trailer and isinstance(pdf.trailer, dict):
                        root = pdf.trailer.get('/Root', {})
                        if isinstance(root, dict):
                            open_action = root.get('/OpenAction')
                            if isinstance(open_action, dict) and open_action.get('/S') == '/JavaScript':
                                found_patterns.append('pdf_open_action')
                                scan_details.append("PDF has JavaScript open action")
                except Exception:
                    pass
                
                # Extract and scan text content from all pages
                for page in pdf.pages:
                    try:
                        text = page.extract_text()
                        if text:
                            # Define truly suspicious patterns for PDFs
                            pdf_suspicious_patterns = {
                                'shellcode': r'\x90{20,}|\xCC{20,}',  # NOP sleds or INT3 instructions
                                'powershell_encoded': r'powershell.*-enc\s+[A-Za-z0-9+/=]+',
                                'suspicious_commands': r'(cmd\.exe|powershell\.exe|wscript\.exe|mshta\.exe).*\/c\s+',
                                'reverse_shell': r'(sh\s*-i\s*>[>&]\s*/dev/tcp/[0-9.]+/\d+|nc\s+-e\s+/bin/[a-z]+|bash\s*-i\s*>[>&]\s*/dev/tcp/)'
                            }
                            
                            # Only check for truly suspicious patterns in PDF text
                            for pattern_name, pattern in pdf_suspicious_patterns.items():
                                if re.search(pattern, text, re.IGNORECASE):
                                    found_patterns.append(f"pdf_{pattern_name}")
                                    scan_details.append(f"Found {pattern_name} pattern in PDF text")
                    except Exception as e:
                        # Only log the first error to avoid spam
                        if not any("Error extracting text from PDF page" in d for d in scan_details):
                            scan_details.append(f"Error extracting text from PDF page: {str(e)}")
            
            except Exception as e:
                scan_details.append(f"Error scanning PDF: {str(e)}")
        
        # Determine verdict based on findings
        if found_patterns:
            if any(p in ['shellcode', 'powershell_encoded', 'suspicious_commands', 'reverse_shell', 'pdf_suspicious_js', 'pdf_open_action'] for p in found_patterns):
                verdict = "Malicious"
            elif len(found_patterns) > 2:
                verdict = "Suspicious"
            else:
                verdict = "Potentially Suspicious"
        else:
            verdict = "Clean"
        
        return (f"{verdict} (Enhanced Scan)", found_patterns, file_hash, scan_details)
            
    except Exception as e:
        return ("Error", [], None, [f"Error: {str(e)}"])

# Process Monitoring
def get_processes():
    processes = []
    # Get CPU percent for all processes at once
    cpu_percents = {}
    for proc in psutil.process_iter(['pid']):
        try:
            cpu_percents[proc.pid] = proc.cpu_percent()
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    
    # Small delay to get accurate CPU percentages
    time.sleep(0.1)
    
    for proc in psutil.process_iter(['pid', 'name', 'cmdline', 'exe', 'memory_percent', 'create_time', 'num_threads', 'ppid']):
        try:
            # Skip System Idle Process and Registry processes
            if proc.info['name'].lower() in ['system idle process', 'registry']:
                continue
                
            cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ""
            
            # Check if the process is a Windows system process
            is_system_process = proc.info['name'].lower() in (p.lower() for p in WINDOWS_SYSTEM_PROCESSES)
            
            # Only check for suspicious keywords if it's not a system process
            is_suspicious = False
            fix_suggestion = ""
            if not is_system_process:
                is_suspicious = any(keyword in cmdline.lower() for keyword in SUSPICIOUS_KEYWORDS)
                if is_suspicious:
                    for keyword in SUSPICIOUS_KEYWORDS:
                        if keyword in cmdline.lower():
                            fix_suggestion = KEYWORD_FIX_SUGGESTIONS.get(keyword, "Investigate this process.")
                            break
            
            # Get process metrics
            cpu_percent = cpu_percents.get(proc.pid, 0) / psutil.cpu_count()
            memory_percent = proc.info.get('memory_percent', 0)
            create_time = datetime.fromtimestamp(proc.info.get('create_time', 0)).strftime('%Y-%m-%d %H:%M:%S')
            num_threads = proc.info.get('num_threads', 0)
            ppid = proc.info.get('ppid', 0)
            
            processes.append({
                'pid': proc.info['pid'],
                'name': proc.info['name'],
                'cmdline': cmdline,
                'suspicious': is_suspicious,
                'system_process': is_system_process,
                'fix_suggestion': fix_suggestion,
                'cpu_percent': cpu_percent,
                'memory_percent': memory_percent,
                'create_time': create_time,
                'num_threads': num_threads,
                'ppid': ppid
            })
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    return processes

def is_private_ip(ip):
    return ip.startswith('10.') or ip.startswith('192.168.') or ip.startswith('172.')

# Add after SUSPICIOUS_IP_RANGES
PROJECT_WHITELIST = {
    'processes': [
        # Development tools
        'python.exe',
        'flask.exe',
        'node.exe',
        'npm.exe',
        'git.exe',
        'vscode.exe',
        'code.exe',
        'powershell.exe',
        'cmd.exe',
        # Windows System Processes
        'svchost.exe',
        'System',
        'services.exe',
        'lsass.exe',
        'winlogon.exe',
        'csrss.exe',
        'spoolsv.exe',
        'wininit.exe',
        'explorer.exe',
        'RuntimeBroker.exe',
        'SearchApp.exe',
        'SearchIndexer.exe',
        'MsMpEng.exe',  # Windows Defender
        'NVDisplay.Container.exe',  # NVIDIA
        'SecurityHealthService.exe',
        'fontdrvhost.exe',
        'dwm.exe',  # Desktop Window Manager
        'WmiPrvSE.exe',  # WMI Provider
        'dllhost.exe',
        'ctfmon.exe',
        'taskhostw.exe',
        'sihost.exe',
        'SearchProtocolHost.exe',
        'SearchFilterHost.exe',
        'WindowsDefender.exe',
        'browser_broker.exe',
        'MicrosoftEdge.exe',
        'msedge.exe',
        'chrome.exe',
        'firefox.exe',
        'WUDFHost.exe',  # Windows Driver Foundation Host
    ],
    'ports': [
        # Common development ports
        5000,  # Flask default port
        3000,  # Node.js default port
        8000,  # Common development port
        8080,  # Common development port
        22,    # SSH
        443,   # HTTPS
        80,    # HTTP
    ],
    'ip_ranges': [
        '127.0.0.1',  # Localhost
        'localhost',
    ],
    'directories': [
        'C:\\Users\\Pixel\\Documents\\ProXDefend',  # Project directory
        'C:\\Users\\Pixel\\Documents\\ProXDefend\\uploads',
        'C:\\Users\\Pixel\\Documents\\ProXDefend\\cache',
    ]
}

def is_suspicious_connection(conn):
    """Enhanced function to detect suspicious network connections."""
    suspicious_reasons = []
    
    try:
        # Check remote address and port
        if conn.raddr:
            remote_port = conn.raddr.port
            remote_ip = conn.raddr.ip
            
            # Check for suspicious ports (excluding whitelisted ports)
            if remote_port in SUSPICIOUS_PORTS and remote_port not in PROJECT_WHITELIST['ports']:
                suspicious_reasons.append(f"Connection to suspicious port {remote_port}")
            
            # Check for private IP ranges (excluding whitelisted IPs)
            if any(ipaddress.ip_address(remote_ip) in ipaddress.ip_network(range) for range in SUSPICIOUS_IP_RANGES):
                if remote_ip not in PROJECT_WHITELIST['ip_ranges']:
                    suspicious_reasons.append(f"Connection to private IP range: {remote_ip}")
            
            # Check for localhost connections (excluding whitelisted)
            if remote_ip == '127.0.0.1' and remote_ip not in PROJECT_WHITELIST['ip_ranges']:
                suspicious_reasons.append("Connection to localhost")
        
        # Check local port
        if conn.laddr:
            local_port = conn.laddr.port
            if local_port in SUSPICIOUS_PORTS and local_port not in PROJECT_WHITELIST['ports']:
                suspicious_reasons.append(f"Listening on suspicious port {local_port}")
        
        # Check connection status
        if conn.status == 'LISTEN' and conn.laddr and conn.laddr.port in SUSPICIOUS_PORTS and conn.laddr.port not in PROJECT_WHITELIST['ports']:
            suspicious_reasons.append(f"Listening on suspicious port {conn.laddr.port}")
        
        # Check for process information
        if conn.pid:
            try:
                process = psutil.Process(conn.pid)
                process_name = process.name().lower()
                
                # Skip whitelisted processes
                if process_name in [p.lower() for p in PROJECT_WHITELIST['processes']]:
                    return False, []
                
                # Check for suspicious process names
                if any(keyword in process_name for keyword in SUSPICIOUS_KEYWORDS):
                    suspicious_reasons.append(f"Suspicious process name: {process_name}")
                
                # Check for unusual number of connections
                process_connections = len(process.connections())
                if process_connections > 10:  # Threshold for multiple connections
                    suspicious_reasons.append(f"Process has {process_connections} active connections")
                
                # Check for unusual process location
                if process.exe():
                    exe_path = process.exe().lower()
                    if not any(loc.lower() in exe_path for loc in PROJECT_WHITELIST['directories']):
                        suspicious_reasons.append(f"Process running from unusual location: {exe_path}")
            
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                suspicious_reasons.append("Unable to access process information")
        
        return len(suspicious_reasons) > 0, suspicious_reasons
    
    except Exception as e:
        return False, [f"Error checking connection: {str(e)}"]

def get_network_connections():
    connections = []
    for conn in psutil.net_connections(kind='inet'):
        try:
            # Get connection details
            laddr = f"{conn.laddr.ip}:{conn.laddr.port}" if conn.laddr else "N/A"
            raddr = f"{conn.raddr.ip}:{conn.raddr.port}" if conn.raddr else "N/A"
            status = conn.status
            pid = conn.pid
            
            # Get process name for all connections
            process_name = "N/A"
            if pid:
                try:
                    process = psutil.Process(pid)
                    process_name = process.name()
                except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                    pass
            
            # Check if the connection is suspicious
            is_suspicious, suspicious_reasons = is_suspicious_connection(conn)
            fix_suggestion = ""
            
            if is_suspicious:
                # Generate appropriate fix suggestion based on the reasons
                if any("suspicious port" in reason.lower() for reason in suspicious_reasons):
                    fix_suggestion = "Consider blocking this port in your firewall."
                elif any("private IP" in reason.lower() for reason in suspicious_reasons):
                    fix_suggestion = "Investigate internal network communication."
                elif any("unusual location" in reason.lower() for reason in suspicious_reasons):
                    fix_suggestion = "Verify the legitimacy of this process."
                else:
                    fix_suggestion = "Investigate this connection for potential security risks."
            
            connections.append({
                'pid': pid,
                'process_name': process_name,
                'laddr': laddr,
                'raddr': raddr,
                'status': status,
                'suspicious': is_suspicious,
                'process_file_hash': None,
                'process_vt_verdict': "Not Checked",
                'fix_suggestion': fix_suggestion,
                'suspicious_reasons': suspicious_reasons
            })
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    return connections

def get_system_uptime():
    boot_time = psutil.boot_time()
    uptime_seconds = time.time() - boot_time
    uptime_hours = int(uptime_seconds // 3600)
    uptime_days = uptime_hours // 24
    uptime_minutes = int((uptime_seconds % 3600) // 60)
    return f"{uptime_days} days, {uptime_hours % 24} hours, {uptime_minutes} minutes"

# Main Route
@app.route('/', methods=['GET'])
def index():
    return render_template('index.html', system_uptime=get_system_uptime())

@app.route('/processes', methods=['GET'])
def processes():
    # Get processes without VirusTotal checks
    process_list = get_processes()
    system_uptime = get_system_uptime()
    
    return render_template('processes.html', 
                         processes=process_list,
                         system_uptime=system_uptime)

@app.route('/network', methods=['GET'])
def network():
    return render_template('network.html', 
                          connections=get_network_connections(),
                          system_uptime=get_system_uptime())

@app.route('/scanner', methods=['GET', 'POST'])
def scanner():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(url_for('scanner'))
        file = request.files['file']
        if file.filename == '':
            return redirect(url_for('scanner'))
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        file_size = os.path.getsize(filepath) / (1024*1024)  # Convert to MB
        file_type = magic.from_file(filepath, mime=True)
        verdict, found_keywords, file_hash, scan_details = scan_file_enhanced(filepath)
        scan_result = {
            'verdict': verdict,
            'file_type': file_type,
            'file_size': f"{file_size:.2f} MB",
            'found_keywords': found_keywords,
            'scan_details': scan_details,
            'filename': file.filename,
            'file_hash': file_hash
        }
        return render_template('scanner.html', scan_result=scan_result, system_uptime=get_system_uptime())
    return render_template('scanner.html', system_uptime=get_system_uptime())

# Simple Test Route
@app.route('/test')
def test():
    return "Test route"

# Network testing routes
@app.route('/ping-test')
def ping_test():
    """Simple endpoint for ping testing"""
    return "pong"

@app.route('/speed-test')
def speed_test():
    """Endpoint for download speed testing"""
    # Generate a 1MB file for testing
    test_data = b'0' * (1024 * 1024)  # 1MB of zeros
    return test_data

@app.route('/upload-test', methods=['POST'])
def upload_test():
    """Endpoint for upload speed testing"""
    # Just acknowledge the upload
    return "OK"

@app.route('/local-ip')
def get_local_ip():
    """Get the local IP address"""
    try:
        # Create a socket connection to get local IP
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
        return jsonify({'ip': local_ip})
    except Exception as e:
        print(f"Error getting local IP: {e}")
        return jsonify({'ip': 'Not Available'})

def generate_processes_pdf(processes):
    # Create a buffer to store the PDF
    buffer = io.BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30
    )
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        wordWrap='LTR'
    )
    
    # Create the content
    content = []
    
    # Add title
    title = Paragraph(f"Process Monitor Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", title_style)
    content.append(title)
    content.append(Spacer(1, 12))
    
    # Prepare table data
    table_data = [['PID', 'Process Name', 'Command Line', 'Status', 'Fix Suggestion']]
    
    for proc in processes:
        status = "Suspicious" if proc['suspicious'] else "Clean"
        table_data.append([
            str(proc['pid']),
            Paragraph(proc['name'], cell_style),
            Paragraph(proc['cmdline'][:200] + '...' if len(proc['cmdline']) > 200 else proc['cmdline'], cell_style),
            Paragraph(status, cell_style),
            Paragraph(proc['fix_suggestion'] if proc['suspicious'] else '', cell_style)
        ])
    
    # Create table
    table = Table(table_data, colWidths=[0.5*inch, 1.5*inch, 3*inch, 1*inch, 2*inch])
    
    # Add style to table
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ])
    
    # Add row colors for suspicious processes
    for i, proc in enumerate(processes, 1):
        if proc['suspicious']:
            style.add('BACKGROUND', (0, i), (-1, i), colors.pink)
    
    table.setStyle(style)
    content.append(table)
    
    # Build PDF
    doc.build(content)
    
    # Get the value of the buffer
    pdf = buffer.getvalue()
    buffer.close()
    
    return pdf

@app.route('/export_processes_pdf')
def export_processes_pdf():
    processes = get_processes()
    pdf = generate_processes_pdf(processes)
    
    # Create a new buffer with the PDF
    buffer = io.BytesIO(pdf)
    
    # Generate filename with timestamp
    filename = f"process_monitor_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )

@app.route('/kill_process/<int:pid>', methods=['POST'])
def kill_process(pid):
    """Safely terminate a process with security checks."""
    try:
        # Get process information
        process = psutil.Process(pid)
        process_name = process.name().lower()
        
        # Security checks
        if process_name in (p.lower() for p in CRITICAL_SYSTEM_PROCESSES):
            return jsonify({
                "success": False,
                "error": "Cannot terminate critical system process",
                "process_name": process_name
            }), 403
            
        if process.username() != psutil.Process().username():
            return jsonify({
                "success": False,
                "error": "Access denied - insufficient permissions",
                "process_name": process_name
            }), 403

        # Get process details for logging
        process_info = {
            "name": process_name,
            "pid": pid,
            "username": process.username(),
            "create_time": datetime.fromtimestamp(process.create_time()).strftime('%Y-%m-%d %H:%M:%S'),
            "cmd": " ".join(process.cmdline()) if process.cmdline() else ""
        }
        
        # Try to terminate gracefully first
        process.terminate()
        
        try:
            # Wait up to 3 seconds for the process to terminate
            process.wait(timeout=3)
            
            return jsonify({
                "success": True,
                "message": f"Process {process_name} (PID: {pid}) terminated successfully",
                "process_info": process_info
            })
            
        except psutil.TimeoutExpired:
            # If graceful termination fails, force kill
            process.kill()
            
            return jsonify({
                "success": True,
                "message": f"Process {process_name} (PID: {pid}) force killed",
                "process_info": process_info
            })
            
    except psutil.NoSuchProcess:
        return jsonify({
            "success": False,
            "error": "Process not found",
            "pid": pid
        }), 404
        
    except psutil.AccessDenied:
        return jsonify({
            "success": False,
            "error": "Access denied - insufficient permissions",
            "pid": pid
        }), 403
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
            "pid": pid
        }), 500

def get_latest_scan_results():
    """Get the latest file scan results from the scan history."""
    try:
        if os.path.exists(SCAN_HISTORY_FILE):
            with open(SCAN_HISTORY_FILE, 'r') as f:
                scan_history = json.load(f)
                if scan_history:
                    # Return the most recent scan result
                    return scan_history[-1]
        return None
    except Exception as e:
        print(f"Error getting latest scan results: {e}")
        return None

@app.route('/process-metrics')
def get_process_metrics():
    """Get CPU and memory usage metrics"""
    try:
        # Get CPU usage
        cpu_percent = psutil.cpu_percent(interval=1)
        
        # Get memory usage
        memory = psutil.virtual_memory()
        memory_percent = memory.percent
        memory_total = memory.total
        memory_used = memory.used
        memory_free = memory.available
        
        return jsonify({
            'cpu_percent': cpu_percent,
            'memory_percent': memory_percent,
            'memory_total': memory_total,
            'memory_used': memory_used,
            'memory_free': memory_free,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
    except Exception as e:
        print(f"Error getting process metrics: {e}")
        return jsonify({
            'cpu_percent': 0,
            'memory_percent': 0,
            'memory_total': 0,
            'memory_used': 0,
            'memory_free': 0,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })

def analyze_system_logs():
    """Analyze system logs for suspicious activities."""
    suspicious_events = []
    try:
        # Windows Event Log analysis
        import win32evtlog
        server = 'localhost'
        logtype = 'Security'
        hand = win32evtlog.OpenEventLog(server, logtype)
        flags = win32evtlog.EVENTLOG_BACKWARDS_READ|win32evtlog.EVENTLOG_SEQUENTIAL_READ
        events = win32evtlog.ReadEventLog(hand, flags, 0)
        
        for event in events:
            if event.EventType in [win32evtlog.EVENTLOG_ERROR_TYPE, win32evtlog.EVENTLOG_WARNING_TYPE]:
                suspicious_events.append({
                    'time': event.TimeGenerated.strftime('%Y-%m-%d %H:%M:%S'),
                    'type': 'Error' if event.EventType == win32evtlog.EVENTLOG_ERROR_TYPE else 'Warning',
                    'source': event.SourceName,
                    'message': event.StringInserts[0] if event.StringInserts else 'No message'
                })
    except Exception as e:
        suspicious_events.append({'error': str(e)})
    return suspicious_events

def monitor_file_changes(directory='.'):
    """Monitor file system changes in real-time."""
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    
    class FileChangeHandler(FileSystemEventHandler):
        def __init__(self):
            self.changes = []
            
        def on_any_event(self, event):
            if not event.is_directory:
                self.changes.append({
                    'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'type': event.event_type,
                    'path': event.src_path
                })
    
    handler = FileChangeHandler()
    observer = Observer()
    observer.schedule(handler, directory, recursive=True)
    observer.start()
    return handler, observer

def check_system_integrity():
    """Perform system integrity checks."""
    integrity_checks = {
        'critical_files': [],
        'suspicious_changes': [],
        'security_issues': []
    }
    
    # Check critical system files
    critical_files = [
        'C:\\Windows\\System32\\kernel32.dll',
        'C:\\Windows\\System32\\user32.dll',
        'C:\\Windows\\System32\\ntdll.dll'
    ]
    
    for file in critical_files:
        if os.path.exists(file):
            file_hash = get_file_hash(file)
            # Here you would compare with known good hashes
            integrity_checks['critical_files'].append({
                'file': file,
                'hash': file_hash,
                'status': 'present'
            })
        else:
            integrity_checks['critical_files'].append({
                'file': file,
                'status': 'missing'
            })
    
    return integrity_checks

def analyze_network_traffic():
    """Analyze network traffic patterns and detect anomalies."""
    traffic_stats = {
        'total_bytes_sent': 0,
        'total_bytes_recv': 0,
        'connections_by_port': {},
        'suspicious_connections': [],
        'traffic_anomalies': []
    }
    
    # Get network connections
    connections = get_network_connections()
    
    # Analyze each connection
    for conn in connections:
        try:
            # Get process info if available
            if conn['pid']:
                process = psutil.Process(conn['pid'])
                process_name = process.name()
                
                # Get connection stats
                if conn['raddr'] != "N/A":
                    remote_port = int(conn['raddr'].split(':')[1])
                    
                    # Track connections by port
                    if remote_port not in traffic_stats['connections_by_port']:
                        traffic_stats['connections_by_port'][remote_port] = {
                            'count': 0,
                            'processes': set()
                        }
                    traffic_stats['connections_by_port'][remote_port]['count'] += 1
                    traffic_stats['connections_by_port'][remote_port]['processes'].add(process_name)
                    
                    # Check for suspicious patterns
                    if remote_port in SUSPICIOUS_PORTS:
                        traffic_stats['suspicious_connections'].append({
                            'pid': conn['pid'],
                            'process': process_name,
                            'local_addr': conn['laddr'],
                            'remote_addr': conn['raddr'],
                            'status': conn['status'],
                            'reason': f"Connection to suspicious port {remote_port}"
                        })
            
            # Get network interface stats
            net_io = psutil.net_io_counters()
            traffic_stats['total_bytes_sent'] = net_io.bytes_sent
            traffic_stats['total_bytes_recv'] = net_io.bytes_recv
            
            # Check for traffic anomalies
            if net_io.bytes_sent > 1000000000 or net_io.bytes_recv > 1000000000:  # 1GB threshold
                traffic_stats['traffic_anomalies'].append({
                    'type': 'high_traffic',
                    'bytes_sent': net_io.bytes_sent,
                    'bytes_recv': net_io.bytes_recv,
                    'threshold': 1000000000
                })
                
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    
    return traffic_stats

def classify_network_connection(conn):
    """Classify a network connection as clean or suspicious."""
    classification = {
        'status': 'clean',
        'confidence': 100,
        'reasons': []
    }
    
    try:
        # Check remote address
        if conn['raddr'] != "N/A":
            remote_port = int(conn['raddr'].split(':')[1])
            
            # Check for suspicious ports
            if remote_port in SUSPICIOUS_PORTS:
                classification['status'] = 'suspicious'
                classification['confidence'] -= 30
                classification['reasons'].append(f"Connection to suspicious port {remote_port}")
            
            # Check for known malicious IPs (you can expand this list)
            remote_ip = conn['raddr'].split(':')[0]
            if remote_ip in KNOWN_MALICIOUS_IPS:
                classification['status'] = 'malicious'
                classification['confidence'] -= 50
                classification['reasons'].append(f"Connection to known malicious IP {remote_ip}")
        
        # Check process behavior
        if conn['pid']:
            process = psutil.Process(conn['pid'])
            
            # Check for suspicious process names
            if any(keyword in process.name().lower() for keyword in SUSPICIOUS_PROCESS_NAMES):
                classification['status'] = 'suspicious'
                classification['confidence'] -= 20
                classification['reasons'].append(f"Suspicious process name: {process.name()}")
            
            # Check for unusual number of connections
            if len(process.connections()) > 10:
                classification['status'] = 'suspicious'
                classification['confidence'] -= 15
                classification['reasons'].append("Unusual number of connections")
        
        # Ensure confidence stays between 0 and 100
        classification['confidence'] = max(0, min(100, classification['confidence']))
        
    except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
        classification['status'] = 'unknown'
        classification['confidence'] = 0
        classification['reasons'].append("Unable to analyze process")
    
    return classification

@app.route('/api/network-traffic', methods=['GET'])
def get_network_traffic_api():
    """API endpoint for network traffic analysis."""
    return jsonify(analyze_network_traffic())

@app.route('/api/connection-classification/<int:pid>', methods=['GET'])
def get_connection_classification_api(pid):
    """API endpoint for connection classification."""
    connections = get_network_connections()
    for conn in connections:
        if conn['pid'] == pid:
            return jsonify(classify_network_connection(conn))
    return jsonify({"error": "Connection not found"}), 404

def detect_malware_behavior():
    """Detect potential malware behavior patterns."""
    suspicious_patterns = []
    
    # Check for suspicious process behavior
    for proc in psutil.process_iter(['pid', 'name', 'cmdline']):
        try:
            process = proc.info
            cmdline = ' '.join(process['cmdline']) if process['cmdline'] else ''
            
            # Check for suspicious command patterns
            for keyword in SUSPICIOUS_KEYWORDS:
                if keyword.lower() in cmdline.lower():
                    suspicious_patterns.append({
                        'type': 'suspicious_command',
                        'process': process['name'],
                        'pid': process['pid'],
                        'pattern': keyword,
                        'command': cmdline
                    })
            
            # Check for suspicious network behavior
            connections = psutil.Process(process['pid']).connections()
            for conn in connections:
                if conn.status == 'LISTEN' and conn.laddr.port in SUSPICIOUS_PORTS:
                    suspicious_patterns.append({
                        'type': 'suspicious_port',
                        'process': process['name'],
                        'pid': process['pid'],
                        'port': conn.laddr.port
                    })
                    
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    
    return suspicious_patterns

def perform_memory_analysis():
    """Analyze system memory for suspicious patterns."""
    memory_analysis = {
        'suspicious_processes': [],
        'high_memory_usage': [],
        'memory_leaks': []
    }
    
    # Check for processes with high memory usage
    for proc in psutil.process_iter(['pid', 'name', 'memory_percent']):
        try:
            process = proc.info
            if process['memory_percent'] > 10:  # Threshold of 10%
                memory_analysis['high_memory_usage'].append({
                    'process': process['name'],
                    'pid': process['pid'],
                    'memory_percent': process['memory_percent']
                })
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    
    return memory_analysis

@app.route('/api/logs', methods=['GET'])
def get_system_logs():
    """API endpoint to get system log analysis."""
    return jsonify(analyze_system_logs())

@app.route('/api/integrity', methods=['GET'])
def get_system_integrity():
    """API endpoint to get system integrity check results."""
    return jsonify(check_system_integrity())

@app.route('/api/malware-detection', methods=['GET'])
def get_malware_detection():
    """API endpoint to get malware behavior detection results."""
    return jsonify(detect_malware_behavior())

@app.route('/api/memory-analysis', methods=['GET'])
def get_memory_analysis():
    """API endpoint to get memory analysis results."""
    return jsonify(perform_memory_analysis())

@app.route('/api/process/top', methods=['GET'])
def get_top_processes():
    """Get top processes by CPU and memory usage."""
    processes = []
    for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent']):
        try:
            process = proc.info
            process['suspicious'] = any(keyword in process['name'].lower() for keyword in SUSPICIOUS_KEYWORDS)
            processes.append(process)
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    
    # Sort by CPU usage and get top 5
    processes.sort(key=lambda x: x['cpu_percent'], reverse=True)
    return jsonify(processes[:5])

@app.route('/api/scan-history', methods=['GET'])
def get_scan_history():
    """Get recent file scan history."""
    scans = cache_manager.get('scan_history', [])
    return jsonify(scans)

@app.route('/api/system-health', methods=['GET'])
def get_system_health():
    """Get overall system health metrics."""
    try:
        cpu_freq = psutil.cpu_freq()
        disk = psutil.disk_usage('/')
        memory = psutil.virtual_memory()
        
        return jsonify({
            'cpu': {
                'percent': psutil.cpu_percent(interval=1),
                'frequency': {
                    'current': cpu_freq.current if cpu_freq else None,
                    'min': cpu_freq.min if cpu_freq else None,
                    'max': cpu_freq.max if cpu_freq else None
                },
                'cores': psutil.cpu_count(logical=False),
                'threads': psutil.cpu_count(logical=True)
            },
            'memory': {
                'total': memory.total,
                'available': memory.available,
                'percent': memory.percent,
                'used': memory.used
            },
            'disk': {
                'total': disk.total,
                'used': disk.used,
                'free': disk.free,
                'percent': disk.percent
            },
            'network': {
                'bytes_sent': psutil.net_io_counters().bytes_sent,
                'bytes_recv': psutil.net_io_counters().bytes_recv,
                'packets_sent': psutil.net_io_counters().packets_sent,
                'packets_recv': psutil.net_io_counters().packets_recv
            },
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
    except Exception as e:
        print(f"Error getting system health: {e}")
        return jsonify({
            'cpu': {
                'percent': 0,
                'frequency': {'current': None, 'min': None, 'max': None},
                'cores': 0,
                'threads': 0
            },
            'memory': {
                'total': 0,
                'available': 0,
                'percent': 0,
                'used': 0
            },
            'disk': {
                'total': 0,
                'used': 0,
                'free': 0,
                'percent': 0
            },
            'network': {
                'bytes_sent': 0,
                'bytes_recv': 0,
                'packets_sent': 0,
                'packets_recv': 0
            },
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })

# Process crash tracking
process_crash_history = {}
process_startup_locations = {}

def track_process_crashes():
    """Track processes that frequently crash and restart."""
    current_processes = {}
    crashed_processes = []
    
    for proc in psutil.process_iter(['pid', 'name', 'create_time']):
        try:
            process_info = proc.info
            process_key = f"{process_info['name']}_{process_info['pid']}"
            
            if process_key in process_crash_history:
                # Check if process has restarted (new create_time)
                if process_info['create_time'] > process_crash_history[process_key]['create_time']:
                    crashed_processes.append({
                        'name': process_info['name'],
                        'pid': process_info['pid'],
                        'crashes': process_crash_history[process_key]['crashes'] + 1,
                        'last_crash': datetime.fromtimestamp(process_info['create_time']).strftime('%Y-%m-%d %H:%M:%S')
                    })
            
            current_processes[process_key] = {
                'create_time': process_info['create_time'],
                'crashes': process_crash_history.get(process_key, {}).get('crashes', 0)
            }
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    
    # Update crash history
    process_crash_history.clear()
    process_crash_history.update(current_processes)
    
    return crashed_processes

def check_startup_locations():
    """Check for processes starting from unusual locations."""
    suspicious_locations = []
    common_locations = [
        'C:\\Windows\\System32',
        'C:\\Program Files',
        'C:\\Program Files (x86)',
        'C:\\Users'
    ]
    
    for proc in psutil.process_iter(['pid', 'name', 'exe']):
        try:
            if proc.info['exe']:
                exe_path = proc.info['exe'].lower()
                is_suspicious = True
                
                # Check if path is in common locations
                for location in common_locations:
                    if location.lower() in exe_path:
                        is_suspicious = False
                        break
                
                if is_suspicious:
                    suspicious_locations.append({
                        'name': proc.info['name'],
                        'pid': proc.info['pid'],
                        'path': proc.info['exe']
                    })
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    
    return suspicious_locations

# Network monitoring enhancements
def get_network_interfaces():
    """Get detailed information about network interfaces including MAC addresses."""
    interfaces = []
    for interface, addrs in psutil.net_if_addrs().items():
        interface_info = {
            'name': interface,
            'addresses': [],
            'mac': None,
            'status': 'up' if interface in psutil.net_if_stats() and psutil.net_if_stats()[interface].isup else 'down'
        }
        
        for addr in addrs:
            if addr.family == psutil.AF_LINK:  # MAC address
                interface_info['mac'] = addr.address
            else:
                interface_info['addresses'].append({
                    'address': addr.address,
                    'netmask': addr.netmask,
                    'family': 'IPv4' if addr.family == socket.AF_INET else 'IPv6'
                })
        
        interfaces.append(interface_info)
    
    return interfaces

# Socket tracking
previous_sockets = set()

def track_socket_changes():
    """Track changes in open sockets between snapshots."""
    current_sockets = set()
    new_sockets = []
    closed_sockets = []
    
    for conn in psutil.net_connections(kind='inet'):
        try:
            # Handle local address
            laddr = f"{conn.laddr[0]}:{conn.laddr[1]}" if conn.laddr else "N/A"
            # Handle remote address
            raddr = f"{conn.raddr[0]}:{conn.raddr[1]}" if conn.raddr else "N/A"
            
            socket_key = f"{laddr}-{raddr}-{conn.status}"
            current_sockets.add(socket_key)
            
            if socket_key not in previous_sockets:
                new_sockets.append({
                    'local_addr': laddr,
                    'remote_addr': raddr,
                    'status': conn.status,
                    'pid': conn.pid
                })
        except (psutil.NoSuchProcess, psutil.AccessDenied, AttributeError):
            continue
    
    # Find closed sockets
    for socket_key in previous_sockets:
        if socket_key not in current_sockets:
            closed_sockets.append(socket_key)
    
    # Update previous sockets
    previous_sockets.clear()
    previous_sockets.update(current_sockets)
    
    return {
        'new_sockets': new_sockets,
        'closed_sockets': closed_sockets
    }

# Traffic monitoring
def detect_traffic_anomalies():
    """Detect unusual spikes in network traffic."""
    current_io = psutil.net_io_counters()
    traffic_stats = {
        'bytes_sent': current_io.bytes_sent,
        'bytes_recv': current_io.bytes_recv,
        'packets_sent': current_io.packets_sent,
        'packets_recv': current_io.packets_recv,
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    
    # Calculate traffic rates (bytes per second)
    if hasattr(detect_traffic_anomalies, 'last_stats'):
        time_diff = (datetime.now() - datetime.strptime(detect_traffic_anomalies.last_stats['timestamp'], '%Y-%m-%d %H:%M:%S')).total_seconds()
        if time_diff > 0:
            traffic_stats['send_rate'] = (current_io.bytes_sent - detect_traffic_anomalies.last_stats['bytes_sent']) / time_diff
            traffic_stats['recv_rate'] = (current_io.bytes_recv - detect_traffic_anomalies.last_stats['bytes_recv']) / time_diff
            
            # Define thresholds (e.g., 1MB/s)
            threshold = 1024 * 1024
            if traffic_stats['send_rate'] > threshold or traffic_stats['recv_rate'] > threshold:
                traffic_stats['anomaly'] = True
                traffic_stats['anomaly_type'] = 'High traffic rate detected'
    
    detect_traffic_anomalies.last_stats = traffic_stats
    return traffic_stats

# Directory monitoring with optimization
directory_state = {
    'last_check': 0,
    'previous_files': {},
    'monitored_dirs': set(['uploads', 'cache']),  # Only monitor specific directories
    'excluded_patterns': [
        r'\.git/',
        r'\.vscode/',
        r'__pycache__/',
        r'\.pyc$',
        r'\.pyo$',
        r'\.pyd$',
        r'\.so$',
        r'\.dll$',
        r'\.exe$',
        r'\.pdf$',
        r'\.docx?$',
        r'\.xlsx?$',
        r'\.pptx?$',
        r'\.txt$',
        r'\.log$',
        r'\.json$',
        r'\.xml$',
        r'\.html$',
        r'\.css$',
        r'\.js$',
        r'\.jpg$',
        r'\.jpeg$',
        r'\.png$',
        r'\.gif$',
        r'\.ico$',
        r'\.svg$',
        r'\.woff$',
        r'\.woff2$',
        r'\.ttf$',
        r'\.eot$',
        r'\.otf$',
        r'\.mp3$',
        r'\.mp4$',
        r'\.avi$',
        r'\.mov$',
        r'\.wmv$',
        r'\.flv$',
        r'\.mkv$',
        r'\.zip$',
        r'\.rar$',
        r'\.7z$',
        r'\.tar$',
        r'\.gz$',
        r'\.bz2$',
        r'\.xz$'
    ]
}

def should_monitor_file(file_path):
    """Check if a file should be monitored based on exclusion patterns."""
    for pattern in directory_state['excluded_patterns']:
        if re.search(pattern, file_path, re.IGNORECASE):
            return False
    return True

def track_directory_changes():
    """Track changes in directory contents with optimized monitoring."""
    current_time = time.time()
    
    # Only check every 30 seconds
    if current_time - directory_state['last_check'] < 30:
        return {
            'new_files': [],
            'modified_files': []
        }
    
    current_files = {}
    new_files = []
    modified_files = []
    
    # Only scan specific directories
    for monitored_dir in directory_state['monitored_dirs']:
        if not os.path.exists(monitored_dir):
            continue
            
        for root, _, files in os.walk(monitored_dir):
            for file in files:
                file_path = os.path.join(root, file)
                
                # Skip files that match exclusion patterns
                if not should_monitor_file(file_path):
                    continue
                
                try:
                    # Get file stats
                    stats = os.stat(file_path)
                    current_files[file_path] = {
                        'size': stats.st_size,
                        'mtime': stats.st_mtime
                    }
                    
                    # Check if file is new or modified
                    if file_path not in directory_state['previous_files']:
                        new_files.append(file_path)
                    elif current_files[file_path]['mtime'] > directory_state['previous_files'][file_path]['mtime']:
                        modified_files.append(file_path)
                except (OSError, IOError) as e:
                    print(f"Error accessing file {file_path}: {e}")
                    continue
    
    # Update state
    directory_state['previous_files'] = current_files
    directory_state['last_check'] = current_time
    
    return {
        'new_files': new_files,
        'modified_files': modified_files
    }

# Add new API endpoints
@app.route('/api/process-crashes')
def get_process_crashes():
    """Get information about frequently crashing processes."""
    return jsonify(track_process_crashes())

@app.route('/api/startup-locations')
def get_startup_locations():
    """Get information about processes starting from unusual locations."""
    return jsonify(check_startup_locations())

@app.route('/api/network-interfaces')
def get_network_interfaces_info():
    """Get detailed information about network interfaces."""
    return jsonify(get_network_interfaces())

@app.route('/api/socket-changes')
def get_socket_changes():
    """Get information about socket changes."""
    return jsonify(track_socket_changes())

@app.route('/api/traffic-anomalies')
def get_traffic_anomalies():
    """Get information about network traffic anomalies."""
    return jsonify(detect_traffic_anomalies())

@app.route('/api/directory-changes')
def get_directory_changes():
    """Get information about directory changes."""
    return jsonify(track_directory_changes())

def generate_network_pdf(connections, network_interfaces, socket_changes, traffic_anomalies, directory_changes):
    """Generate and download a PDF report of network monitoring data."""
    try:
        print("Starting PDF generation...")
        
        # Gather all the data
        connections = get_network_connections()
        network_interfaces = get_network_interfaces()
        socket_changes = track_socket_changes()
        traffic_anomalies = detect_traffic_anomalies()
        directory_changes = track_directory_changes()
        
        print("Data gathered successfully")
        
        # Create a buffer to store the PDF
        buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=20
        )
        
        content = []
        
        # Add title
        content.append(Paragraph("Network Monitoring Report", title_style))
        content.append(Spacer(1, 20))
        
        # Add timestamp
        content.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        content.append(Spacer(1, 20))
        
        # Network Connections
        content.append(Paragraph("Network Connections", subtitle_style))
        if connections:
            table_data = [['PID', 'Process', 'Local Address', 'Remote Address', 'Status']]
            for conn in connections:
                table_data.append([
                    str(conn['pid']),
                    conn['process_name'],
                    conn['laddr'],
                    conn['raddr'],
                    conn['status']
                ])
            
            table = Table(table_data, colWidths=[1*inch, 2*inch, 2*inch, 2*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(table)
            content.append(Spacer(1, 20))
        
        # Network Interfaces
        content.append(Paragraph("Network Interfaces", subtitle_style))
        if network_interfaces:
            table_data = [['Interface', 'MAC Address', 'Status', 'IP Addresses']]
            for interface in network_interfaces:
                table_data.append([
                    interface['name'],
                    interface['mac'] or 'N/A',
                    interface['status'],
                    '\n'.join([f"{addr['address']} ({addr['family']})" for addr in interface['addresses']])
                ])
            
            table = Table(table_data, colWidths=[2*inch, 2*inch, 1*inch, 3*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(table)
        
        print("Building PDF...")
        # Build the PDF
        doc.build(content)
        
        # Get the value of the buffer
        pdf = buffer.getvalue()
        buffer.close()
        
        print("PDF generated successfully")
        
        # Create a new buffer for sending
        output = io.BytesIO(pdf)
        output.seek(0)
        
        # Generate filename with timestamp
        filename = f"network_monitor_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        print(f"Sending PDF file: {filename}")
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        print("Traceback:")
        traceback.print_exc()
        return jsonify({"error": "Failed to generate PDF report"}), 500

@app.route('/export_network_pdf')
def export_network_pdf():
    try:
        print("Starting PDF generation...")
        
        # Create a buffer to store the PDF
        buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30
        )
        
        content = []
        
        # Add title
        content.append(Paragraph("Network Monitoring Report", title_style))
        content.append(Spacer(1, 20))
        
        # Add timestamp
        content.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        content.append(Spacer(1, 20))
        
        # Get network data
        connections = get_network_connections()
        network_interfaces = get_network_interfaces()
        
        # Network Connections
        content.append(Paragraph("Network Connections", styles['Heading2']))
        if connections:
            table_data = [['PID', 'Process', 'Local Address', 'Remote Address', 'Status']]
            for conn in connections:
                table_data.append([
                    str(conn['pid']),
                    conn['process_name'],
                    conn['laddr'],
                    conn['raddr'],
                    conn['status']
                ])
            
            table = Table(table_data, colWidths=[1*inch, 2*inch, 2*inch, 2*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(table)
            content.append(Spacer(1, 20))
        
        # Network Interfaces
        content.append(Paragraph("Network Interfaces", styles['Heading2']))
        if network_interfaces:
            table_data = [['Interface', 'MAC Address', 'Status', 'IP Addresses']]
            for interface in network_interfaces:
                table_data.append([
                    interface['name'],
                    interface['mac'] or 'N/A',
                    interface['status'],
                    '\n'.join([f"{addr['address']} ({addr['family']})" for addr in interface['addresses']])
                ])
            
            table = Table(table_data, colWidths=[2*inch, 2*inch, 1*inch, 3*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(table)
        
        print("Building PDF...")
        # Build the PDF
        doc.build(content)
        
        # Get the value of the buffer
        pdf = buffer.getvalue()
        buffer.close()
        
        print("PDF generated successfully")
        
        # Create a new buffer for sending
        output = io.BytesIO(pdf)
        output.seek(0)
        
        # Generate filename with timestamp
        filename = f"network_monitor_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        print(f"Sending PDF file: {filename}")
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        print("Traceback:")
        traceback.print_exc()
        return jsonify({"error": "Failed to generate PDF report"}), 500

def detect_process_anomalies():
    """Detect anomalous process behavior patterns."""
    anomalies = []
    
    # Get process tree
    process_tree = get_process_tree()
    
    # Analyze process patterns
    for pid, proc in process_tree.items():
        try:
            process = psutil.Process(pid)
            
            # Skip anomaly detection for Windows system processes
            if process.name().lower() in (p.lower() for p in WINDOWS_SYSTEM_PROCESSES):
                continue
            
            # Check for unusual CPU spikes
            cpu_percent = process.cpu_percent(interval=0.1)
            if cpu_percent > 90:  # Threshold for CPU spike
                anomalies.append({
                    'type': 'cpu_spike',
                    'pid': pid,
                    'name': proc['name'],
                    'value': cpu_percent,
                    'threshold': 90
                })
            
            # Check for unusual memory usage
            memory_percent = process.memory_percent()
            if memory_percent > 50:  # Threshold for memory usage
                anomalies.append({
                    'type': 'high_memory',
                    'pid': pid,
                    'name': proc['name'],
                    'value': memory_percent,
                    'threshold': 50
                })
            
            # Check for unusual number of threads
            num_threads = process.num_threads()
            if num_threads > 100:  # Threshold for thread count
                anomalies.append({
                    'type': 'high_threads',
                    'pid': pid,
                    'name': proc['name'],
                    'value': num_threads,
                    'threshold': 100
                })
            
            # Check for unusual child processes
            if len(proc['children']) > 10:  # Threshold for child processes
                anomalies.append({
                    'type': 'many_children',
                    'pid': pid,
                    'name': proc['name'],
                    'value': len(proc['children']),
                    'threshold': 10
                })
            
        except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
            continue
    
    return anomalies

@app.route('/api/process-anomalies', methods=['GET'])
def get_process_anomalies_api():
    """API endpoint for process anomalies."""
    return jsonify(detect_process_anomalies())

# Network connection action endpoints
@app.route('/api/connection-details/<int:pid>')
def get_connection_details(pid):
    """Get detailed information about a specific network connection."""
    try:
        # Get process information
        process = psutil.Process(pid)
        connections = process.connections()
        
        if not connections:
            return jsonify({'error': 'No active connections found for this process'}), 404
            
        # Get the first connection (most relevant)
        conn = connections[0]
        
        return jsonify({
            'process_path': process.exe(),
            'start_time': datetime.fromtimestamp(process.create_time()).strftime('%Y-%m-%d %H:%M:%S'),
            'protocol': 'TCP' if conn.type == socket.SOCK_STREAM else 'UDP',
            'bytes_sent': process.io_counters().write_bytes,
            'bytes_recv': process.io_counters().read_bytes
        })
    except (psutil.NoSuchProcess, psutil.AccessDenied) as e:
        return jsonify({'error': str(e)}), 404

@app.route('/api/connection-traffic/<int:pid>')
def get_connection_traffic(pid):
    """Get real-time traffic information for a specific connection."""
    try:
        process = psutil.Process(pid)
        # Get current IO counters
        io = process.io_counters()
        
        # Calculate rates (bytes per second)
        # Note: In a production environment, you'd want to store previous values
        # and calculate actual rates over time
        upload_rate = io.write_bytes / (time.time() - process.create_time())
        download_rate = io.read_bytes / (time.time() - process.create_time())
        
        # Calculate utilization (example thresholds)
        max_rate = 100 * 1024 * 1024  # 100 MB/s as example threshold
        upload_utilization = min((upload_rate / max_rate) * 100, 100)
        download_utilization = min((download_rate / max_rate) * 100, 100)
        
        return jsonify({
            'upload_rate': upload_rate,
            'download_rate': download_rate,
            'upload_utilization': upload_utilization,
            'download_utilization': download_utilization
        })
    except (psutil.NoSuchProcess, psutil.AccessDenied) as e:
        return jsonify({'error': str(e)}), 404

@app.route('/api/block-connection/<int:pid>', methods=['POST'])
def block_connection(pid):
    """Block a network connection using Windows Firewall."""
    try:
        process = psutil.Process(pid)
        exe_path = process.exe()
        
        # Create firewall rule using netsh
        rule_name = f"Block_{os.path.basename(exe_path)}_{pid}"
        command = f'netsh advfirewall firewall add rule name="{rule_name}" dir=out program="{exe_path}" action=block'
        
        result = subprocess.run(command, shell=True, capture_output=True, text=True)
        
        if result.returncode == 0:
            return jsonify({'success': True})
        else:
            return jsonify({
                'success': False,
                'error': result.stderr or 'Failed to create firewall rule'
            })
    except (psutil.NoSuchProcess, psutil.AccessDenied) as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/whitelist-connection/<int:pid>', methods=['POST'])
def whitelist_connection(pid):
    """Whitelist a network connection by removing any blocking rules."""
    try:
        process = psutil.Process(pid)
        exe_path = process.exe()
        
        # Remove any existing firewall rules for this process
        rule_name = f"Block_{os.path.basename(exe_path)}_{pid}"
        command = f'netsh advfirewall firewall delete rule name="{rule_name}"'
        
        result = subprocess.run(command, shell=True, capture_output=True, text=True)
        
        # Create allow rule
        allow_rule_name = f"Allow_{os.path.basename(exe_path)}_{pid}"
        allow_command = f'netsh advfirewall firewall add rule name="{allow_rule_name}" dir=out program="{exe_path}" action=allow'
        
        allow_result = subprocess.run(allow_command, shell=True, capture_output=True, text=True)
        
        if allow_result.returncode == 0:
            return jsonify({'success': True})
        else:
            return jsonify({
                'success': False,
                'error': allow_result.stderr or 'Failed to create allow rule'
            })
    except (psutil.NoSuchProcess, psutil.AccessDenied) as e:
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    app.run(debug=True)
